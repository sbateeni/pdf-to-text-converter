from docx import Document
import os

def save_text_to_word(texts, pdf_filename, search_query=""):
    # Determine the output directory relative to the script file
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(script_dir, 'output')

    # Create the output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Define the path for the Word file
    base_name = os.path.splitext(pdf_filename)[0]
    word_file_path = os.path.join(output_dir, f"{base_name}.docx")

    # Create a new Word document
    doc = Document()
    for i, text in enumerate(texts):
        doc.add_heading(f'Page {i + 1}', level=1)
        if search_query:
            text = highlight_text(text, search_query)  # Apply highlighting if search_query is present
        doc.add_paragraph(text)

    # Save the document to the specified path
    doc.save(word_file_path)
    return word_file_path

def highlight_text(text, search_query):
    """Highlight occurrences of search_query in text with a yellow background."""
    highlighted_text = re.sub(
        re.escape(search_query),
        lambda m: f'<span style="background-color: yellow;">{m.group(0)}</span>',
        text,
        flags=re.IGNORECASE
    )
    return highlighted_text
